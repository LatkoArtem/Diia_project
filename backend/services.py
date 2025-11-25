from docx import Document
import os

def generate_contract_docx(template_path: str, answers: dict, output_path: str):
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    doc = Document(template_path)

    # 1. Параграфи
    for para in doc.paragraphs:
        replace_text_preserving_style(para, answers)

    # 2. Таблиці
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_preserving_style(para, answers)

    doc.save(output_path)
    return output_path

def replace_text_preserving_style(paragraph, answers):
    """
    Розумна заміна:
    1. Перевіряє, чи є ключі в тексті параграфа.
    2. Якщо є - робить заміну в повному тексті.
    3. Видаляє старий вміст параграфа.
    4. Записує новий текст, копіюючи стиль (шрифт, розмір, жирність) з оригінального початку.
    """
    text = paragraph.text
    # Швидка перевірка, чи варто взагалі чіпати цей параграф
    has_changes = False
    for key, value in answers.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in text:
            text = text.replace(placeholder, str(value))
            has_changes = True

    if has_changes:
        # Зберігаємо стиль першого 'run' (фрагмента), якщо він є
        style_run = paragraph.runs[0] if paragraph.runs else None

        # Очищаємо параграф
        paragraph.clear()

        # Створюємо новий run з новим текстом
        new_run = paragraph.add_run(text)

        # Копіюємо основні стилі (можна розширити список)
        if style_run:
            new_run.bold = style_run.bold
            new_run.italic = style_run.italic
            new_run.underline = style_run.underline
            new_run.font.name = style_run.font.name
            new_run.font.size = style_run.font.size
            new_run.font.color.rgb = style_run.font.color.rgb