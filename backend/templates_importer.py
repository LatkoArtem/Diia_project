import os
import re
import json
from groq import Groq
from docx import Document
from dotenv import load_dotenv
import models

# Завантажуємо налаштування
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
MODEL_NAME = "llama-3.1-8b-instant"

def extract_placeholders(text):
    """Знаходить {{KEY}} у тексті"""
    return [m.strip() for m in re.findall(r"\{\{(.*?)\}\}", text)]

def ask_llm_about_slot(client, key: str):
    """
    Бере системний КЛЮЧ, очищує його перед LLM, і генерує
    ідеальне запитання для користувача.
    """
    key = key.strip()

    processed_key = key.replace("_", " ").title()

    prompt = f"""
Ти — помічник-лінгвіст. Твоє завдання — перетворити англійську фразу (PHRASE) на граматично бездоганне запитання для користувача українською мовою.

--- СУВОРІ ПРАВИЛА ---
1. Мова: Українська.
2. Граматика: Завжди використовуй правильні відмінки (напр., "Director Full Name" -> "Введіть повне ім'я директора").
3. Початок: Завжди починай з дієслова з великої літери ("Введіть", "Вкажіть").
4. Кінець: НІКОЛИ не став крапку чи знак питання в кінці.
5. Формат: Тільки JSON з одним ключем "question".

--- ІДЕАЛЬНІ ПРИКЛАДИ ---
PHRASE: "Director Full Name"
JSON: {{"question": "Введіть повне ім'я директора"}}

PHRASE: "Customer Iban"
JSON: {{"question": "Вкажіть поточний рахунок замовника"}}

PHRASE: "Contract Date"
JSON: {{"question": "Вкажіть дату укладання договору"}}

PHRASE: "Payment Amount"
JSON: {{"question": "Вкажіть суму платежу"}}

--- ТВОЄ ЗАВДАННЯ ---
Дотримуючись усіх правил, створи JSON для цієї фрази:
PHRASE: "{processed_key}"
JSON:
"""
    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            response_format={"type": "json_object"}
        )
        raw = response.choices[0].message.content.strip()
        data = json.loads(raw)

        if "question" not in data or not data["question"]:
             data = {"question": f"Введіть {processed_key.lower()}"}

    except Exception as e:
        print(f"LLM error for {key}: {e}")
        data = {"question": f"Введіть {processed_key.lower()}"}

    return data

def generate_json_schema_for_docx(docx_path):
    """Генерує повну JSON схему для файлу"""
    if not GROQ_API_KEY:
        print("⚠️ SKIPPING AI GENERATION: No GROQ_API_KEY found in .env")
        return {}

    client = Groq(api_key=GROQ_API_KEY)
    doc = Document(docx_path)
    all_keys = set()

    for para in doc.paragraphs:
        all_keys.update(extract_placeholders(para.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    all_keys.update(extract_placeholders(para.text))

    print(f"🤖 AI аналізує {os.path.basename(docx_path)}... Знайдено {len(all_keys)} полів.")

    slots = {}
    for key in all_keys:
        # 'key' тут - це ОРИГІНАЛЬНИЙ ключ (напр. "DIRECTOR_FULL_NAME")
        info = ask_llm_about_slot(client, key)

        # 4. Зберігаємо результат під ОРИГІНАЛЬНИМ ключем
        slots[key] = info

    return slots

def run_auto_import(db):
    """Головна функція: шукає файли і додає в БД"""
    folder = "storage/templates"
    if not os.path.exists(folder):
        os.makedirs(folder)

    # === 1. СЛОВНИК НАЗВ ===
    # Тут ви прописуєте, як називати кожен файл для користувача
    TEMPLATE_NAMES = {
        "nadannya_poslug": "Договір надання послуг (ФОП)",
        "nda": "Угода про нерозголошення (NDA)",
        "rent_apartment": "Договір оренди квартири",
        # Додавайте нові файли сюди
    }

    # Перебираємо всі .docx файли в папці
    for filename in os.listdir(folder):
        if not filename.endswith(".docx"):
            continue

        # Код шаблону = назва файлу без .docx (напр. "nadannya_poslug")
        code = os.path.splitext(filename)[0]

        # Перевіряємо, чи вже є такий шаблон в базі
        existing = db.query(models.ContractTemplate).filter_by(code=code).first()
        if existing:
            continue

        print(f"🆕 Новий файл знайдено: {filename}. Імпортуємо...")

        full_path = os.path.join(folder, filename)

        # 1. Генеруємо питання через AI
        json_schema = generate_json_schema_for_docx(full_path)

        # === 2. ВИЗНАЧАЄМО НАЗВУ ===
        # Якщо код є в нашому словнику - беремо українську назву.
        # Якщо немає - просто робимо красиву назву з файлу.
        nice_name = TEMPLATE_NAMES.get(code, code.replace("_", " ").title())

        # 3. Записуємо в Базу Даних
        new_template = models.ContractTemplate(
            name=nice_name,  # <--- Ось тут тепер буде українська назва
            code=code,
            json_schema=json_schema,
            docx_path=full_path
        )
        db.add(new_template)
        db.commit()
        print(f"✅ Успішно додано: {nice_name}")